"""
Generate the (condensed) expert validation Word document for the NMT thesis.
Run: .venv-nmt/Scripts/python -m scripts.generate_expert_validation_doc
Output: reports/05_nmt/evaluation_xl/validacion_experto_nmt.docx

Condensed ~2-page version aligned with the university template: objective,
expert info, results, IOVs, a short 1-5 rubric, final verdict + 2 open
questions, signature and acknowledgement.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = Path("reports/05_nmt/evaluation_xl/validacion_experto_nmt.docx")


# ── helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(13 if level == 1 else 11)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    run.bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_paragraph(doc, text="", size=10, space_after=4):
    p = doc.add_paragraph()
    if text:
        run = p.add_run(text)
        run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_table_header_row(table, headers, bg="1F497D", text_color="FFFFFF"):
    row = table.rows[0]
    tc = RGBColor(int(text_color[0:2], 16), int(text_color[2:4], 16), int(text_color[4:6], 16))
    for i, h in enumerate(headers):
        cell = row.cells[i]
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = tc
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_data_row(row, alt=False, center_cols=None):
    center_cols = center_cols or []
    for i, cell in enumerate(row.cells):
        set_cell_bg(cell, "EEF3F9" if alt else "FFFFFF")
        for p in cell.paragraphs:
            if i in center_cols:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(9)


def info_table(doc, rows, label_w=Cm(5.5), value_w=Cm(10)):
    """Two-column label/value table with shaded label column."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(rows):
        row = table.rows[i]
        set_cell_bg(row.cells[0], "D6E4F0")
        r = row.cells[0].paragraphs[0].add_run(label)
        r.bold = True
        r.font.size = Pt(9)
        row.cells[1].paragraphs[0].add_run(value).font.size = Pt(9)
        row.cells[0].width = label_w
        row.cells[1].width = value_w
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return table


def answer_lines(doc, n=3):
    for _ in range(n):
        lp = doc.add_paragraph("_" * 110)
        lp.paragraph_format.space_before = Pt(0)
        lp.paragraph_format.space_after = Pt(1)
        lp.runs[0].font.size = Pt(9)
        lp.runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)


# ── document ──────────────────────────────────────────────────────────────────

def build():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    # ── Title ─────────────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run("DOCUMENTO PARA VALIDACION CON EXPERTOS")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Experto en Traduccion Automatica y Evaluacion de Sistemas NMT")
    rs.font.size = Pt(10)
    rs.bold = True
    rs.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    sub.paragraph_format.space_after = Pt(8)

    info_table(doc, [
        ("Titulo del proyecto",
         "Desarrollo de un Sistema de Traduccion Automatica Neuronal Bidireccional Shiwilu-Castellano"),
        ("Estudiante", "Fabian Fernando Prado Infanson (20206480)"),
        ("Asesor", "Erasmo Gomez Perez"),
        ("Fecha de validacion", "____ / ____ / 2026"),
    ])

    # ── 1. Objetivo ───────────────────────────────────────────────────────────
    add_heading(doc, "1. Objetivo")
    add_paragraph(
        doc,
        "Obtener la validacion de un experto respecto a la metodologia de evaluacion, la "
        "interpretacion de los resultados y la seleccion del modelo final del proyecto de tesis. "
        "En particular, se solicita revisar la pertinencia de las metricas utilizadas (BLEU y "
        "chrF++), la lectura diferenciada por direccion de traduccion, el analisis cualitativo y "
        "estadistico, y la evidencia que sustenta la eleccion del modelo campeon.",
    )

    # ── 2. Informacion del experto ────────────────────────────────────────────
    add_heading(doc, "2. Informacion del experto")
    info_table(doc, [
        ("Nombre", ""),
        ("Cargo / Especialidad", ""),
        ("Institucion", ""),
        ("Correo electronico", ""),
    ])

    # ── 3. Resultados obtenidos ───────────────────────────────────────────────
    add_heading(doc, "3. Resultados obtenidos")
    add_paragraph(
        doc,
        "La evaluacion se hizo por direccion: BLEU para shiwilu->castellano y chrF++ para "
        "castellano->shiwilu (chrF++ como criterio global). El modelo base partia en el umbral de "
        "ruido (BLEU 9.7 / chrF++ 4.6); el modelo campeon (v2.1b, LoRA+) lo supera con holgura "
        "(BLEU 24.5 / chrF++ promedio 44.99, IC 95 % [43.17, 46.96]).",
    )

    # Distribucion por bandas
    bands = doc.add_table(rows=4, cols=3)
    bands.style = "Table Grid"
    add_table_header_row(bands, ["Banda de calidad", "Shiwilu->Castellano (BLEU)", "Castellano->Shiwilu (chrF++)"])
    bands_data = [
        ("Alta (BLEU >=20 / chrF++ >=40)", "70 %", "55 %"),
        ("Media (BLEU 10-20 / chrF++ 20-40)", "13 %", "33 %"),
        ("Ruido (BLEU <10 / chrF++ <20)", "17 %", "11 %"),
    ]
    for i, (band, a, b) in enumerate(bands_data):
        row = bands.rows[i + 1]
        row.cells[0].paragraphs[0].add_run(band).font.size = Pt(9)
        row.cells[1].paragraphs[0].add_run(a).font.size = Pt(9)
        row.cells[2].paragraphs[0].add_run(b).font.size = Pt(9)
        style_data_row(row, alt=(i % 2 == 1), center_cols=[1, 2])
    bands.columns[0].width = Cm(6)
    bands.columns[1].width = Cm(5)
    bands.columns[2].width = Cm(5)

    add_paragraph(doc, "", space_after=2)

    # Comparacion finalistas
    finals = doc.add_table(rows=3, cols=3)
    finals.style = "Table Grid"
    add_table_header_row(finals, ["Modelo finalista", "chrF++ promedio", "IC bootstrap 95 %"])
    finals_data = [
        ("v2.1b - LoRA+ (campeon)", "44.99", "[43.17, 46.96]"),
        ("v2.1  - DoRA + LoRA+", "44.30", "[42.49, 46.12]"),
    ]
    for i, vals in enumerate(finals_data):
        row = finals.rows[i + 1]
        for j, val in enumerate(vals):
            row.cells[j].paragraphs[0].add_run(val).font.size = Pt(9)
        style_data_row(row, alt=(i % 2 == 1), center_cols=[1, 2])
    finals.columns[0].width = Cm(6)
    finals.columns[1].width = Cm(5)
    finals.columns[2].width = Cm(5)
    add_paragraph(
        doc,
        "Los intervalos se solapan: la diferencia no es estadisticamente significativa. Se "
        "selecciona v2.1b por parsimonia (mismo nivel con un adaptador mas simple).",
        space_after=4,
    )

    # ── 4. IOV propuestos ─────────────────────────────────────────────────────
    add_heading(doc, "4. Indicadores Objetivamente Verificables (IOV)")
    iov = doc.add_table(rows=5, cols=3)
    iov.style = "Table Grid"
    add_table_header_row(iov, ["IOV", "Definicion operacional", "Resultado obtenido"])
    iov_data = [
        ("Protocolo de evaluacion reproducible",
         "Modelos evaluados en condiciones equivalentes y documentadas.",
         "Mismo conjunto de prueba y metricas definidas previamente para todos los modelos."),
        ("Reporte comparativo con metricas automaticas",
         "Resultados que permiten comparar objetivamente los modelos.",
         "BLEU y chrF++ reportados para todas las variantes experimentales."),
        ("Analisis cualitativo complementario",
         "Puntajes interpretados con ejemplos reales de traduccion.",
         "Traducciones clasificadas por bandas de calidad y revisadas manualmente."),
        ("Seleccion justificada con incertidumbre",
         "Eleccion del modelo final sustentada en evidencia y estabilidad estadistica.",
         "Intervalos de confianza bootstrap y analisis de significancia practica."),
    ]
    for i, (a, b, c) in enumerate(iov_data):
        row = iov.rows[i + 1]
        row.cells[0].paragraphs[0].add_run(a).font.size = Pt(9)
        row.cells[1].paragraphs[0].add_run(b).font.size = Pt(9)
        row.cells[2].paragraphs[0].add_run(c).font.size = Pt(9)
        style_data_row(row, alt=(i % 2 == 1))
    iov.columns[0].width = Cm(4)
    iov.columns[1].width = Cm(6)
    iov.columns[2].width = Cm(6)

    # ── 5. Evaluacion del experto (rubrica) ───────────────────────────────────
    add_heading(doc, "5. Evaluacion del experto")
    add_paragraph(
        doc,
        "Asigne un puntaje de 1 a 5 (1 = no adecuado, 3 = aceptable, 5 = muy adecuado).",
        space_after=4,
    )
    rubric = doc.add_table(rows=6, cols=3)
    rubric.style = "Table Grid"
    add_table_header_row(rubric, ["Criterio", "Puntaje (1-5)", "Comentario"])
    rubric_data = [
        "Pertinencia de las metricas (BLEU shw->spa, chrF++ spa->shw) para este par de lenguas.",
        "La interpretacion de los resultados es razonable y consistente con el estado del arte.",
        "El analisis cualitativo complementa adecuadamente las metricas automaticas.",
        "Los IC bootstrap y la comparacion entre modelos estan correctamente planteados.",
        "Coherencia entre objetivos, metodologia, resultados y conclusiones.",
    ]
    for i, crit in enumerate(rubric_data):
        row = rubric.rows[i + 1]
        row.cells[0].paragraphs[0].add_run(crit).font.size = Pt(9)
        row.cells[1].paragraphs[0].add_run("").font.size = Pt(9)
        row.cells[2].paragraphs[0].add_run("").font.size = Pt(9)
        style_data_row(row, alt=(i % 2 == 1), center_cols=[1])
    rubric.columns[0].width = Cm(9)
    rubric.columns[1].width = Cm(2.5)
    rubric.columns[2].width = Cm(4.5)

    # ── 6. Dictamen final + preguntas abiertas ────────────────────────────────
    add_heading(doc, "6. Dictamen y comentarios")
    add_paragraph(doc, "Marque la alternativa que refleje su opinion general:", space_after=3)
    for opt in [
        "Validado sin observaciones relevantes.",
        "Validado con observaciones menores.",
        "Validado con observaciones que deben atenderse antes de la version final.",
        "No validado en su estado actual.",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run("☐  " + opt)
        r.font.size = Pt(11)

    add_paragraph(doc, "", space_after=2)
    q1 = doc.add_paragraph()
    q1.add_run("1. ¿Que analisis adicionales recomendaria antes de la version final de la tesis?").bold = True
    q1.runs[0].font.size = Pt(10)
    q1.paragraph_format.space_after = Pt(2)
    answer_lines(doc, 2)

    add_paragraph(doc, "", space_after=2)
    q2 = doc.add_paragraph()
    q2.add_run("2. ¿Tiene observaciones sobre la metodologia para comparar y seleccionar el modelo final?").bold = True
    q2.runs[0].font.size = Pt(10)
    q2.paragraph_format.space_after = Pt(2)
    answer_lines(doc, 2)

    # ── 7. Firma + agradecimiento ─────────────────────────────────────────────
    add_heading(doc, "7. Firma del experto")
    info_table(doc, [
        ("Nombre", ""),
        ("Firma", ""),
        ("Fecha", "____ / ____ / 2026"),
    ])
    add_paragraph(doc, "", space_after=4)
    add_paragraph(
        doc,
        "Agradezco sinceramente el tiempo dedicado a la revision de este trabajo. Sus observaciones "
        "contribuiran a fortalecer la metodologia de evaluacion y la calidad final de la investigacion.",
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Documento guardado en: {OUTPUT}")


if __name__ == "__main__":
    build()
